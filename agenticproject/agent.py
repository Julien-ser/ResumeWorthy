import os
import io
import json
import re
from typing import Literal, Dict, Any, List
from dotenv import load_dotenv
import streamlit as st
import pandas as pd
from langchain_core.tools import tool
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from langchain_core.messages import HumanMessage
from ddgs import DDGS

import pypdf
import docx

load_dotenv()
def extract_text(file) -> str:
    if not file:
        return ""
    name = file.name.lower()
    if name.endswith(".txt"):
        return file.read().decode("utf-8", errors="ignore")
    if name.endswith(".pdf"):
        pdf = pypdf.PdfReader(io.BytesIO(file.read()))
        return "\n".join((p.extract_text() or "") for p in pdf.pages)
    if name.endswith(".docx"):
        d = docx.Document(io.BytesIO(file.read()))
        return "\n".join(p.text for p in d.paragraphs)
    return ""
def md_to_docx(md_text: str) -> bytes:
    doc = docx.Document()
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 3)
            doc.add_heading(line.lstrip("#").strip(), level=level)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def normalize_jobs(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    normed = []
    for it in items:
        if not isinstance(it, dict):
            continue
        # case-insensitive keys
        lower_map = {str(k).strip().lower(): it[k] for k in it.keys()}
        company = str(lower_map.get("company", "") or "").strip()
        title = str(lower_map.get("title", "") or "").strip()
        location = str(lower_map.get("location", "") or "").strip()
        link = str(lower_map.get("link", "") or "").strip()
        why_fit = str(lower_map.get("why_fit", lower_map.get("good match", "")) or "").strip()
        if not link:
            continue
        normed.append({
            "company": company or "—",
            "title": title or "—",
            "location": location or "—",
            "link": link,
            "Good Match": "Yes" if why_fit else "—",
        })
    return normed[:5]
def extract_jobs_from_text(text: str) -> List[Dict[str, Any]]:
    if not text:
        return []
    pattern = r"<JOBS>\s*(?:```[\w-]*\s*)?(\[.*?\])\s*(?:```)?\s*</JOBS>"
    m = re.search(pattern, text, flags=re.S | re.I)
    if not m:
        return []
    raw = m.group(1).strip().strip("`").strip()
    try:
        obj = json.loads(raw)
        return obj if isinstance(obj, list) else []
    except Exception:
        try:
            salvaged = re.sub(r"(?<!\\)'", '"', raw)
            obj = json.loads(salvaged)
            return obj if isinstance(obj, list) else []
        except Exception:
            st.session_state.last_error = f"JSON parse failed: {raw[:1200]}"
            return []

if "jobs_df" not in st.session_state:
    st.session_state.jobs_df = None
if "cover_doc" not in st.session_state:
    st.session_state.cover_doc = None
if "last_error" not in st.session_state:
    st.session_state.last_error = ""
if "raw_final" not in st.session_state:
    st.session_state.raw_final = ""

st.set_page_config(page_title="Job Application Assistant", page_icon=" ", layout="wide")
st.title("💼 Job Application Assistant")
c0, c1, c2 = st.columns([2, 1, 1])
with c0:
    uploaded = st.file_uploader("Upload your resume (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"])
with c1:
    target_title = st.text_input("Target title", "Intern AI Engineer")
with c2:
    target_location = st.text_input("Target location(s)", "Toronto OR Remote")
skills_hint = st.text_area(
    "Add/override skills (optional)",
    "",
    placeholder="Python, PyTorch, LLMs, RAG, Azure, vLLM, FastAPI",
)


@tool
def internet_search(
    query: str,
    max_results: int = 5,
    topic: Literal["general", "news", "finance"] = "general",
    include_raw_content: bool = False,
) -> List[Dict[str, Any]]:
    """Search the internet using DuckDuckGo to find relevant information and job postings. No API key required."""
    try:
        ddgs = DDGS()
        results = ddgs.text(query, max_results=max_results)
        formatted = []
        for r in results:
            formatted.append({
                "title": r.get("title", ""),
                "url": r.get("href", ""),
                "content": r.get("body", "")
            })
        return formatted
    except Exception as e:
        raise RuntimeError(f"Search failed: {e}")


INSTRUCTIONS = (
    "You are a job application assistant. Your task:\n"
    "1) Search for exactly 5 CURRENT job postings matching the user's target title, locations, and skills.\n"
    "2) Return ONLY valid JSON in this exact format:\n"
    "<JOBS>\n"
    "[{\"company\":\"Company Name\",\"title\":\"Job Title\",\"location\":\"City, State\",\"link\":\"https://actual-job-url.com\",\"Good Match\":\"one sentence why this matches\"},"
    " ... 5 total jobs]\n"
    "</JOBS>\n"
    "3) All links MUST be real, working URLs to the actual job posting or company career page.\n"
    "4) Do NOT invent or fake links. Use real sources: company.com/careers, linkedin.com, indeed.com, Lever, Greenhouse platforms.\n"
    "5) Each link must be clickable and go to the actual job posting.\n"
    "CRITICAL: Ensure all links are complete URLs starting with https:// and are real, verifiable job posting pages."
)
JOB_SEARCH_PROMPT = (
    "Search and select 5 real postings that match the user's title, locations, and skills. "
    "Output ONLY this block format (no extra text before/after the wrapper):\n"
    "<JOBS>\n"
    "[{\"company\":\"...\",\"title\":\"...\",\"location\":\"...\",\"link\":\"https://...\",\"Good Match\":\"one sentence\"},"
    " {\"company\":\"...\",\"title\":\"...\",\"location\":\"...\",\"link\":\"https://...\",\"Good Match\":\"one sentence\"},"
    " {\"company\":\"...\",\"title\":\"...\",\"location\":\"...\",\"link\":\"https://...\",\"Good Match\":\"one sentence\"},"
    " {\"company\":\"...\",\"title\":\"...\",\"location\":\"...\",\"link\":\"https://...\",\"Good Match\":\"one sentence\"},"
    " {\"company\":\"...\",\"title\":\"...\",\"location\":\"...\",\"link\":\"https://...\",\"Good Match\":\"one sentence\"}]"
    "\n</JOBS>"
)
COVER_LETTER_PROMPT = (
    "For each job in the found list, write a subject line and a concise cover letter (≤150 words) that ties the user's skills/resume to the role. "
    "Append to cover_letters.md under a heading per job. Keep writing tight and specific."
)
def build_agent():
    api_key = os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        st.error("Please set OPENROUTER_API_KEY in your environment.")
        st.stop()
    llm = ChatOpenAI(
        model="arcee-ai/trinity-large-preview:free",
        api_key=api_key,
        base_url="https://openrouter.ai/api/v1",
        temperature=0.2
    )
    return llm
def make_task_prompt(resume_text: str, skills_hint: str, title: str, location: str) -> str:
    skills = skills_hint.strip()
    skill_line = f" Prioritize these skills: {skills}." if skills else ""
    return (
        f"Target title: {title}\n"
        f"Target location(s): {location}\n"
        f"{skill_line}\n\n"
        f"RESUME RAW TEXT:\n{resume_text[:8000]}"
    )


resume_text = extract_text(uploaded) if uploaded else ""
run_clicked = st.button("Run", type="primary", disabled=not uploaded)
if run_clicked:
    st.session_state.last_error = ""
    st.session_state.raw_final = ""
    try:
        if not os.environ.get("OPENROUTER_API_KEY"):
            st.error("OPENROUTER_API_KEY not set.")
            st.stop()
        agent = build_agent()
        task = make_task_prompt(resume_text, skills_hint, target_title, target_location)
        
        # Actually search for jobs using the search tool
        with st.spinner("Searching for real job postings..."):
            # Search with more specific queries to find actual job postings
            search_queries = [
                f'"{target_title}" jobs {target_location} site:linkedin.com',
                f'"{target_title}" {target_location} career site:greenhouse.io OR site:lever.co',
                f'{target_title} remote {target_location} hiring',
                f'{target_title} openings {target_location}',
            ]
            
            raw_jobs = []
            for search_query in search_queries:
                if len(raw_jobs) >= 5:
                    break
                try:
                    search_results = internet_search.invoke({"query": search_query, "max_results": 5})
                    
                    for result in search_results:
                        if len(raw_jobs) >= 5:
                            break
                        url = result.get("url", "")
                        title = result.get("title", "")
                        content = result.get("content", "")
                        
                        # Filter for actual job posting URLs (not just boards)
                        if not url.startswith("http"):
                            continue
                        
                        # Extract company from title or content
                        company = "Company"
                        if " at " in title:
                            company = title.split(" at ")[-1].strip()
                        elif " - " in title:
                            # Usually "Company - Job Title"
                            parts = title.split(" - ")
                            if len(parts) > 1:
                                company = parts[0].strip()
                        elif "|" in title:
                            # Sometimes "Company | Job Title"
                            company = title.split("|")[0].strip()
                        else:
                            # Try to extract from URL (e.g., linkedin.com/jobs/view/123456-ai-ops-intern-at-klue)
                            if "linkedin.com" in url and "-at-" in url:
                                url_part = url.split("-at-")[-1].split("/")[0].split("?")[0]
                                if url_part:
                                    company = url_part.replace("-", " ").title()
                            # Or from domain (e.g., klue.com)
                            elif url.count("/") > 2:
                                domain = url.split("//")[1].split("/")[0]
                                if domain.startswith("www."):
                                    domain = domain[4:]
                                company_name = domain.split(".")[0]
                                if len(company_name) > 1 and company_name.isalpha():
                                    company = company_name.upper()
                        
                        # Skip if URL is just a board landing page
                        if any(board in url.lower() for board in ["linkedin.com/feed", "indeed.com/jobs", "glassdoor.com/jobs"]):
                            continue
                        
                        job = {
                            "company": company,
                            "title": target_title,
                            "location": target_location.split(" OR ")[0],
                            "link": url,
                            "Good Match": content[:80] if content else "Job posting"
                        }
                        raw_jobs.append(job)
                    
                except Exception as e:
                    st.warning(f"Query '{search_query}' failed: {e}")
                    continue
            
            if raw_jobs:
                st.info(f"Found {len(raw_jobs)} job postings.")
            else:
                st.warning("No specific job postings found. Try different search criteria.")
        
        # Generate cover letters for each job found
        if raw_jobs:
            with st.spinner("Generating personalized cover letters..."):
                cover_md = "# Cover Letters\n\n"
                for i, job in enumerate(raw_jobs, 1):
                    if not isinstance(job, dict):
                        continue
                    company = job.get("company", "Company")
                    title = job.get("title", "Position")
                    location = job.get("location", "")
                    
                    letter_prompt = (
                        f"Write a professional, personalized cover letter (100-150 words) for this position.\n\n"
                        f"Company: {company}\n"
                        f"Position: {title}\n"
                        f"Location: {location}\n\n"
                        f"Resume highlights:\n{resume_text[:1500]}\n\n"
                        f"Key skills to emphasize: {skills_hint}\n\n"
                        f"Write a compelling cover letter that explains why this candidate is a great fit for this specific role."
                    )
                    letter_response = agent.invoke([HumanMessage(content=letter_prompt)])
                    letter_text = letter_response.content if hasattr(letter_response, 'content') else str(letter_response)
                    
                    cover_md += f"## {i}. {company} - {title}\n\n{letter_text}\n\n---\n\n"
        else:
            cover_md = ""
        
        st.session_state.cover_doc = md_to_docx(cover_md) if cover_md else None
        
        jobs_list = normalize_jobs(raw_jobs)
        st.session_state.jobs_df = pd.DataFrame(jobs_list) if jobs_list else None
        
        if raw_jobs:
            st.success(f"Done! Found {len(raw_jobs)} jobs and generated cover letters.")
        else:
            st.warning("No jobs found. Try adjusting your search criteria.")
    except Exception as e:
        st.session_state.last_error = str(e)
        st.error(f"Error: {e}")

st.header("Jobs")
if st.session_state.jobs_df is None or st.session_state.jobs_df.empty:
    st.write("No jobs to show yet.")
else:
    df = st.session_state.jobs_df.copy()
    def as_link(u: str) -> str:
        u = u if isinstance(u, str) else ""
        return f'<a href="{u}" target="_blank">Apply</a>' if u else "—"
    if "link" in df.columns:
        df["link"] = df["link"].apply(as_link)
    cols = [c for c in ["company", "title", "location", "link", "Good Match"] if c in df.columns]
    df = df[cols]
    st.write(df.to_html(escape=False, index=False), unsafe_allow_html=True)
st.header("Download")
if st.session_state.cover_doc:
    st.download_button(
        "Download cover_letters.docx",
        data=st.session_state.cover_doc,
        file_name="cover_letters.docx",
      mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="dl_cover_letters",
    )
else:
    st.caption("Cover letters not produced yet.")